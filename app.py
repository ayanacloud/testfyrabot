import json
import os
import sys
import boto3
import streamlit as st
import docx
import io
import numpy as np
import pickle
from PyPDF2 import PdfReader
from transformers import AutoTokenizer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.chat_models import BedrockChat
from langchain_community.embeddings import BedrockEmbeddings
from langchain.llms.bedrock import Bedrock
from langchain.vectorstores import FAISS
from langchain.schema import Document
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA

# Fetch credentials and region from secrets
aws_access_key_id = st.secrets["aws"]["aws_access_key_id"]
aws_secret_access_key = st.secrets["aws"]["aws_secret_access_key"]
region = st.secrets["aws"]["region"]

# Initialize AWS Clients with credentials and region
s3 = boto3.client(
    "s3",
    aws_access_key_id=aws_access_key_id,
    aws_secret_access_key=aws_secret_access_key,
    region_name=region
)
bedrock = boto3.client(
    "bedrock-runtime",
    aws_access_key_id=aws_access_key_id,
    aws_secret_access_key=aws_secret_access_key,
    region_name=region
)

st.write(st.secrets["aws"])

bedrock_embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v2:0", client=bedrock)

# S3 Bucket Info
bucket_name = "testfyrabot"
prefix = "botfiles/" 

# Load Tokenizer
tokenizer = AutoTokenizer.from_pretrained("gpt2")


### **🔹 Utility Functions**
def list_s3_files():
    """List all files in S3 under a given prefix."""
    response = s3.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
    return [obj["Key"] for obj in response.get("Contents", []) if not obj["Key"].endswith("/")]

def read_text_from_s3(file_key):
    """Read text files (TXT, PDF, DOCX) from S3."""
    response = s3.get_object(Bucket=bucket_name, Key=file_key)
    file_content = response["Body"].read()
    
    # Auto-detect file type
    if file_key.lower().endswith(".pdf"):
        pdf_reader = PdfReader(io.BytesIO(file_content))
        return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
    
    elif file_key.lower().endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs])
    
    else:  # Assume it's a text file
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("ISO-8859-1", errors="ignore")

def data_ingestion():
    """Fetch documents from S3 and return LangChain Document objects."""
    documents = []
    for file_key in list_s3_files():
        text_content = read_text_from_s3(file_key)
        if text_content:
            documents.append(Document(page_content=text_content, metadata={"source": file_key}))
    return documents


### **🔹 Optimized Chunking Function**
def split_documents(docs, chunk_size=512, chunk_overlap=50):
    """Splits documents into smaller chunks for FAISS embedding."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,  
        chunk_overlap=chunk_overlap,  
        length_function=len,  
        separators=["\n\n", "\n", " "]  # Prioritizing meaningful breaks
    )
    return text_splitter.split_documents(docs)


### **🔹 FAISS Vector Store Handling**
def upload_faiss_to_s3(faiss_local_path):
    """Upload FAISS index files to S3."""
    s3.upload_file(f"{faiss_local_path}/index.faiss", bucket_name, "faiss_store/index.faiss")
    s3.upload_file(f"{faiss_local_path}/index.pkl", bucket_name, "faiss_store/index.pkl")
    print("✅ FAISS index uploaded to S3.")

def check_faiss_exists_in_s3():
    """Check if FAISS index exists in S3."""
    try:
        s3.head_object(Bucket=bucket_name, Key="faiss_store/index.faiss")
        return True
    except s3.exceptions.ClientError:
        return False

def load_faiss_from_s3():
    """Load FAISS index from S3 if it exists, else return None."""
    faiss_local_path = "faiss_index"
    os.makedirs(faiss_local_path, exist_ok=True)

    if check_faiss_exists_in_s3():
        print("📥 Downloading FAISS index from S3...")
        s3.download_file(bucket_name, "faiss_store/index.faiss", f"{faiss_local_path}/index.faiss")
        s3.download_file(bucket_name, "faiss_store/index.pkl", f"{faiss_local_path}/index.pkl")
        return FAISS.load_local(faiss_local_path, bedrock_embeddings, allow_dangerous_deserialization=True)
    else:
        print("❌ FAISS index not found in S3.")
        return None

def get_vector_store(docs):
    """Create a FAISS vector store from documents."""
    print("📌 Creating vector store...")
    
    # Apply optimized chunking
    chunked_docs = split_documents(docs)

    # Create FAISS index
    vectorstore_faiss = FAISS.from_documents(chunked_docs, bedrock_embeddings)
    faiss_local_path = "faiss_index"
    vectorstore_faiss.save_local(faiss_local_path)
    print("✅ FAISS index saved locally.")

    return vectorstore_faiss


### **🔹 LLM Handling**
def get_claude_llm():
    """Initialize Claude 3 Sonnet LLM."""
    return BedrockChat(model_id="anthropic.claude-3-sonnet-20240229-v1:0", client=bedrock)

def get_llama2_llm():
    """Initialize LLaMA3-70B model."""
    return Bedrock(model_id="meta.llama3-70b-instruct-v1:0", client=bedrock, model_kwargs={'max_gen_len': 512})


### **🔹 Retrieval + Prompt Handling**
prompt_template = """
Human: You are the Testfyra Bot, part of our team. Use the following pieces of context to provide a 
concise answer to the question at the end but use at least summarize with 
50 words with detailed explanations. If you don't know the answer, 
let us know that you don't know and suggest contacting us through our website. 
Please mention that they can reach out to us by clicking the 'Contact Us' page at www.testfyra.com.
<context>
{context}
</context>

Question: {question}

Assistant (Testfyra Bot): We're here to help! Here's what we can share:
"""

PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

def get_response_llm(llm, vectorstore_faiss, query):
    """Retrieve context from FAISS and generate response."""
    qa = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore_faiss.as_retriever(search_type="similarity", search_kwargs={"k": 3}),
        return_source_documents=True,
        chain_type_kwargs={"prompt": PROMPT}
    )
    return qa({"query": query})['result']


### **🔹 Streamlit UI**
def main():
    st.set_page_config("Chat PDF")
    st.header("Testfyra Bot")

    user_question = st.text_input("Ask a Question & Choose from the models provided:")
    
    with st.sidebar:
        st.title("Vector Store Management")
        if st.button("Update Vectors"):
            with st.spinner("Processing..."):
                docs = data_ingestion()
                get_vector_store(docs)
                upload_faiss_to_s3("faiss_index")
                st.success("✅ FAISS updated!")
    if st.button("Claude Output"):
        with st.spinner("Processing..."):
            faiss_index = load_faiss_from_s3()
            if faiss_index:
                llm = get_claude_llm()
                st.write(get_response_llm(llm, faiss_index, user_question))
                st.success("✅ Done!")

    if st.button("Llama2 Output"):
        with st.spinner("Processing..."):
            faiss_index = load_faiss_from_s3()
            if faiss_index:
                llm = get_llama2_llm()
                st.write(get_response_llm(llm, faiss_index, user_question))
                st.success("✅ Done!")

if __name__ == "__main__":
    main()
